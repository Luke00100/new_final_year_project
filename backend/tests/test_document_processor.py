"""
Unit Tests for Document Processor Module (document_processor.py)
Tests text extraction from PDF, DOCX, and TXT files, plus routing logic.
"""

import pytest
import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

# ── Adjust this import to match your project structure ──
from services.document_processor import (
    extract_text_from_pdf,
    extract_text_from_docx,
    extract_text_from_txt,
    extract_text,
)


# ============================================================
# HELPER FUNCTIONS — create test files in memory
# ============================================================

def create_test_pdf(text: str) -> bytes:
    """Create a simple single-page PDF containing the given text."""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    # Write text near the top of the page
    y_position = 750
    for line in text.split("\n"):
        c.drawString(72, y_position, line)
        y_position -= 15
    c.save()
    return buffer.getvalue()


def create_test_docx(text: str) -> bytes:
    """Create a simple DOCX containing the given text as paragraphs."""
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_test_txt(text: str) -> bytes:
    """Create UTF-8 encoded bytes from a string."""
    return text.encode("utf-8")


# ============================================================
# TXT EXTRACTION TESTS
# ============================================================

class TestExtractTextFromTxt:
    """Tests for plain text extraction."""

    def test_basic_text(self):
        """Should extract basic text content."""
        content = "Hello, this is a test document."
        result = extract_text_from_txt(create_test_txt(content))
        assert result == content

    def test_multiline_text(self):
        """Should preserve multiline content."""
        content = "Line one\nLine two\nLine three"
        result = extract_text_from_txt(create_test_txt(content))
        assert "Line one" in result
        assert "Line two" in result
        assert "Line three" in result

    def test_strips_whitespace(self):
        """Should strip leading and trailing whitespace."""
        content = "   padded text   "
        result = extract_text_from_txt(create_test_txt(content))
        assert result == "padded text"

    def test_empty_file(self):
        """Should return empty string for empty file."""
        result = extract_text_from_txt(b"")
        assert result == ""

    def test_special_characters(self):
        """Should handle special characters correctly."""
        content = "Price: £100 — 50% off! © 2024"
        result = extract_text_from_txt(create_test_txt(content))
        assert result == content


# ============================================================
# DOCX EXTRACTION TESTS
# ============================================================

class TestExtractTextFromDocx:
    """Tests for DOCX text extraction."""

    def test_basic_docx(self):
        """Should extract text from a simple DOCX."""
        file_bytes = create_test_docx("This is a test paragraph.")
        result = extract_text_from_docx(file_bytes)
        assert "This is a test paragraph." in result

    def test_multiline_docx(self):
        """Should extract multiple paragraphs."""
        text = "First paragraph\nSecond paragraph\nThird paragraph"
        file_bytes = create_test_docx(text)
        result = extract_text_from_docx(file_bytes)
        assert "First paragraph" in result
        assert "Second paragraph" in result
        assert "Third paragraph" in result

    def test_empty_docx(self):
        """Should handle a DOCX with no text content."""
        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        result = extract_text_from_docx(buffer.getvalue())
        assert result == ""

    def test_docx_returns_string(self):
        """Should return a string type."""
        file_bytes = create_test_docx("Test content")
        result = extract_text_from_docx(file_bytes)
        assert isinstance(result, str)


# ============================================================
# PDF EXTRACTION TESTS
# ============================================================

class TestExtractTextFromPdf:
    """Tests for PDF text extraction."""

    def test_basic_pdf(self):
        """Should extract text from a simple PDF."""
        file_bytes = create_test_pdf("This is a test PDF document.")
        result = extract_text_from_pdf(file_bytes)
        assert "This is a test PDF document." in result

    def test_multiline_pdf(self):
        """Should extract multiple lines from a PDF."""
        file_bytes = create_test_pdf("First line\nSecond line")
        result = extract_text_from_pdf(file_bytes)
        assert "First line" in result
        assert "Second line" in result

    def test_pdf_returns_string(self):
        """Should return a string type."""
        file_bytes = create_test_pdf("Test content")
        result = extract_text_from_pdf(file_bytes)
        assert isinstance(result, str)

    def test_empty_pdf(self):
        """Should handle a PDF with no text without raising an error."""
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        c.save()
        result = extract_text_from_pdf(buffer.getvalue())
        assert isinstance(result, str)


# ============================================================
# ROUTING / extract_text() TESTS
# ============================================================

class TestExtractTextRouting:
    """Tests for the file-type routing function."""

    def test_routes_txt(self):
        """Should route .txt files to the TXT extractor."""
        content = "Plain text content"
        result = extract_text("readme.txt", create_test_txt(content))
        assert result == content

    def test_routes_docx(self):
        """Should route .docx files to the DOCX extractor."""
        file_bytes = create_test_docx("Word document content")
        result = extract_text("report.docx", file_bytes)
        assert "Word document content" in result

    def test_routes_pdf(self):
        """Should route .pdf files to the PDF extractor."""
        file_bytes = create_test_pdf("PDF content here")
        result = extract_text("document.pdf", file_bytes)
        assert "PDF content here" in result

    def test_case_insensitive_extension(self):
        """Should handle uppercase file extensions."""
        content = "Uppercase extension"
        result = extract_text("FILE.TXT", create_test_txt(content))
        assert result == content

    def test_mixed_case_extension(self):
        """Should handle mixed-case extensions like .Docx."""
        file_bytes = create_test_docx("Mixed case test")
        result = extract_text("report.Docx", file_bytes)
        assert "Mixed case test" in result

    def test_unsupported_extension_raises_error(self):
        """Should raise ValueError for unsupported file types."""
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text("image.png", b"fake image data")

    def test_unsupported_csv_raises_error(self):
        """Should raise ValueError for CSV files."""
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text("data.csv", b"col1,col2\nval1,val2")

    def test_no_extension_raises_error(self):
        """Should raise ValueError for files with no extension."""
        with pytest.raises(ValueError):
            extract_text("noextension", b"some content")